"""Document-to-text extraction for ``read_file``.

Supports Jupyter notebooks (.ipynb), Word documents (.docx via python-docx),
and Excel workbooks (.xlsx via openpyxl). Malformed documents raise
:class:`ExtractionError`; callers can then fall back to normal text/binary
handling.
"""

from __future__ import annotations

import json
from pathlib import Path

__all__ = [
    "EXTRACTABLE_EXTENSIONS",
    "ExtractionError",
    "extract_document_text",
    "is_extractable_document",
]

EXTRACTABLE_EXTENSIONS = frozenset({".ipynb", ".docx", ".xlsx"})


class ExtractionError(Exception):
    """Raised when a supported-looking document cannot be rendered as text."""


def _extension(path: str) -> str:
    ext = Path(path).suffix.lower()
    return ext if ext in EXTRACTABLE_EXTENSIONS else ""


def is_extractable_document(path: str) -> bool:
    return bool(_extension(path))


def extract_document_text(path: str) -> str:
    ext = _extension(path)
    if ext == ".ipynb":
        return _extract_notebook(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".xlsx":
        return _extract_xlsx(path)
    raise ExtractionError(f"Unsupported document type: {path!r}")


# ---------------------------------------------------------------------------
# Notebook (.ipynb)
# ---------------------------------------------------------------------------


def _source_text(source) -> str:
    if isinstance(source, str):
        return source
    if isinstance(source, list):
        return "".join(item for item in source if isinstance(item, str))
    return ""


def _extract_notebook(path: str) -> str:
    try:
        with open(path, encoding="utf-8", errors="replace") as fh:
            nb = json.load(fh)
    except (OSError, ValueError, json.JSONDecodeError) as exc:
        raise ExtractionError(f"Not a valid notebook: {exc}") from exc
    if not isinstance(nb, dict):
        raise ExtractionError("Notebook root is not an object")

    cells = nb.get("cells")
    if not isinstance(cells, list):
        cells = [
            cell
            for ws in nb.get("worksheets", [])
            if isinstance(ws, dict)
            for cell in ws.get("cells", [])
        ]
    if not cells:
        raise ExtractionError("Notebook contains no cells")

    counts = {"markdown": 0, "code": 0, "raw": 0}
    labels = {"markdown": "Markdown", "code": "Code", "raw": "Raw"}
    out: list[str] = []
    for cell in cells:
        if not isinstance(cell, dict):
            continue
        typ = cell.get("cell_type")
        if typ not in labels:
            continue
        counts[typ] += 1
        suffix = f" {counts[typ]}" if typ != "raw" else ""
        out.extend(
            (
                f"# ── {labels[typ]} cell{suffix} ──",
                _source_text(cell.get("source", "")).rstrip("\n"),
                "",
            )
        )
    if not out:
        raise ExtractionError("Notebook contains no readable cells")
    return "\n".join(out).rstrip("\n") + "\n"


# ---------------------------------------------------------------------------
# Word (.docx) — python-docx
# ---------------------------------------------------------------------------


def _extract_docx(path: str) -> str:
    try:
        from docx import Document as _Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError:
        raise ExtractionError("python-docx is required for .docx extraction") from None

    try:
        doc = _Document(path)
    except PackageNotFoundError as exc:
        raise ExtractionError(f"Not a valid DOCX: {exc}") from exc
    except (KeyError, AttributeError, OSError) as exc:
        raise ExtractionError(str(exc)) from exc

    lines: list[str] = []

    # Paragraphs (body text)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Tables
    for table in doc.tables:
        rows_text: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if any(r.strip(" |") for r in rows_text):
            lines.append("")
            lines.extend(rows_text)
            lines.append("")

    # Headers and footers from each section
    for section in doc.sections:
        for header in (section.header,):
            if header and not header.is_linked_to_previous:
                for para in header.paragraphs:
                    text = para.text.strip()
                    if text:
                        lines.append(f"[Header] {text}")
        for footer in (section.footer,):
            if footer and not footer.is_linked_to_previous:
                for para in footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        lines.append(f"[Footer] {text}")

    if not any(line.strip() for line in lines):
        raise ExtractionError("DOCX contains no extractable text")
    return "\n".join(lines).rstrip("\n") + "\n"


# ---------------------------------------------------------------------------
# Excel (.xlsx) — openpyxl
# ---------------------------------------------------------------------------


def _extract_xlsx(path: str) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ExtractionError("openpyxl is required for .xlsx extraction") from None

    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except OSError as exc:
        raise ExtractionError(str(exc)) from exc
    except Exception as exc:
        raise ExtractionError(f"Not a valid XLSX: {exc}") from exc

    out: list[str] = []
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            if ws.sheet_state in {"hidden", "veryHidden"}:
                continue

            out.append(f"# ── Sheet: {sheet_name} ──")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                # Filter fully empty rows
                cleaned = [str(v) if v is not None else "" for v in row]
                if not any(cell.strip() for cell in cleaned):
                    continue
                out.append("\t".join(cleaned))
                row_count += 1
            if row_count == 0:
                out.append("(empty)")
            out.append("")
    finally:
        wb.close()

    if not out:
        raise ExtractionError("XLSX has no visible sheets with content")
    return "\n".join(out).rstrip("\n") + "\n"
